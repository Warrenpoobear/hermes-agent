from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE_PATH = Path(
    r"C:\Users\DarrenSchulz\.codex\attachments\36199bb3-58fb-4a96-bf74-b72a0634cf45\pasted-text.txt"
)
OUT_PATH = Path(r"C:\Projects\hermes-agent\artifacts\documents\Investment_Strategy_Statement_Darren_Schulz_2026-06-25.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
GRAY = RGBColor(85, 85, 85)
LIGHT_FILL = "F4F6F9"
RULE = "B8C7D9"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_rule(paragraph, color=RULE, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.start_type = WD_SECTION_START.NEW_PAGE

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet = styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    bullet._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(8)
    bullet.paragraph_format.line_spacing = 1.167

    num = styles["List Number"]
    num.font.name = "Calibri"
    num._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    num._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    num.font.size = Pt(11)
    num.paragraph_format.left_indent = Inches(0.5)
    num.paragraph_format.first_line_indent = Inches(-0.25)
    num.paragraph_format.space_after = Pt(8)
    num.paragraph_format.line_spacing = 1.167


def set_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.text = ""
    p.paragraph_format.space_after = Pt(0)
    left = p.add_run("Investment Strategy Statement")
    set_run_font(left, size=9, color=GRAY, bold=True)
    right = p.add_run("\tDarren Schulz")
    set_run_font(right, size=9, color=GRAY)
    set_paragraph_rule(p, color="D7DBE2", size="4")

    footer = section.footer
    f = footer.paragraphs[0]
    f.text = ""
    f.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = f.add_run("Page ")
    set_run_font(run, size=9, color=GRAY)
    add_page_field(f)


def add_title_block(doc, title, subtitle):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("INVESTMENT STRATEGY STATEMENT")
    set_run_font(run, size=23, color=RGBColor(0, 0, 0), bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("Darren Schulz")
    set_run_font(run, size=14, color=GRAY)

    rows = [
        ("Date", "June 25, 2026"),
        ("Purpose", "Durable household/family-office capital reference and professional mandate."),
        ("Scope", "Strategy/philosophy statement; not a compliance IPS or advice."),
        ("Posture", "Capital preservation first; automation gated by manual validation."),
    ]
    for label, value in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        label_run = p.add_run(f"{label}: ")
        set_run_font(label_run, size=11, color=RGBColor(0, 0, 0), bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=11, color=RGBColor(0, 0, 0))

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(10)
    rule.paragraph_format.space_after = Pt(12)
    set_paragraph_rule(rule, color=RULE, size="8")

    callout = doc.add_paragraph()
    callout.paragraph_format.left_indent = Inches(0.12)
    callout.paragraph_format.right_indent = Inches(0.12)
    callout.paragraph_format.space_before = Pt(0)
    callout.paragraph_format.space_after = Pt(12)
    callout.paragraph_format.line_spacing = 1.167
    shade_paragraph(callout, LIGHT_FILL)
    r = callout.add_run("Reference frame: ")
    set_run_font(r, size=10.5, color=NAVY, bold=True)
    r = callout.add_run(subtitle)
    set_run_font(r, size=10.5, color=RGBColor(0, 0, 0))


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.keep_together = True
    run = p.add_run(text)
    set_run_font(run, size=11, color=RGBColor(0, 0, 0))
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.keep_together = True
    run = p.add_run(text)
    set_run_font(run, size=11, color=RGBColor(0, 0, 0))
    return p


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.keep_together = True
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=11, color=RGBColor(0, 0, 0))
    return p


def add_exclusion_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.167
    shade_paragraph(p, "F2F4F7")
    r = p.add_run("Excluded from the core book: ")
    set_run_font(r, size=10.5, color=NAVY, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=RGBColor(0, 0, 0))


def add_layer_table(doc):
    rows = [
        ("Core", "Quality compounders", "Concentrated pricing-power book."),
        ("Satellite", "Biotech coinvest", "Signal-driven, tax-advantaged, separately governed."),
        ("Architecture", "Family-office / endowment", "Liquidity, pacing, allocation, and spending discipline."),
        ("Mandate", "Real estate", "Digital-first, hospitality-minded operating platform."),
    ]
    doc.add_heading("Architecture at a glance", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = [Inches(1.15), Inches(1.95), Inches(3.40)]
    hdr = table.rows[0].cells
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for cell, text, width in zip(hdr, ["Layer", "Expression", "Role"], widths):
        cell.width = width
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=9.5, color=NAVY, bold=True)
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "E8EEF5")
        tc_pr.append(shd)
    for row in rows:
        cells = table.add_row().cells
        for cell, text, width in zip(cells, row, widths):
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run_font(r, size=9.5, color=RGBColor(0, 0, 0))


def build():
    text = SOURCE_PATH.read_text(encoding="utf-8")
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    title = lines[0]
    subtitle = lines[1]

    doc = Document()
    configure_document(doc)
    set_header_footer(doc)
    add_title_block(doc, title, subtitle)
    add_layer_table(doc)

    section_headings = {
        "Purpose & first principles",
        "Structural exclusions (core book)",
        "Process & governance (the throughline)",
        "Coherence note",
        "Suggested review cadence",
    }

    current_heading = None
    for line in lines[2:]:
        if line in section_headings or line.startswith("Layer "):
            current_heading = line
            level = 1 if line.startswith("Layer ") else 2
            doc.add_heading(line, level=level)
            continue

        if current_heading == "Structural exclusions (core book)":
            add_exclusion_box(doc, line)
        elif current_heading == "Coherence note":
            p = add_body_paragraph(doc, line)
            shade_paragraph(p, "F7F9FC")
            p.paragraph_format.left_indent = Inches(0.12)
            p.paragraph_format.right_indent = Inches(0.12)
        elif current_heading == "Suggested review cadence":
            add_numbered(doc, line)
        else:
            add_bullet(doc, line)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build())
